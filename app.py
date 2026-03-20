import streamlit as st
import fitz
import pytesseract
from PIL import Image, ImageOps, ImageFilter
import io
import re
import gc
import hashlib

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------
# CONFIG STREAMLIT
# ----------------------------

st.set_page_config(layout="wide", page_title="Noticias Yucatán")
st.set_option("client.showErrorDetails", False)

# Si lo necesitas en Windows, descomenta esta línea:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

DEBUG_MODE = True

# ----------------------------
# CONFIG OCR / RENDER
# ----------------------------

OCR_DPI = 200
OCR_LANG = "spa+eng"
MIN_CHARS_TEXTO_EMBEDIDO = 40

# Calidad de render para mostrar/guardar página
DISPLAY_DPI = 150

# Se ve pequeña en pantalla, pero el PNG es de buena calidad
DISPLAY_WIDTH = 220

# ----------------------------
# PALABRAS CLAVE
# ----------------------------

KEYWORDS = {
    "Yucatán": [
        r"\bYucat[aá]n\b"
    ],
    "YUC": [
        r"\bYUC\b"
    ],
    "yucateco": [
        r"\byucateco\b",
        r"\byucateca\b",
        r"\byucatecos\b",
        r"\byucatecas\b"
    ],
    "Mérida": [
        r"\bM[eé]rida\b"
    ],
    "Gobernador": [
        r"\bGobernador\s+de\s+Yucat[aá]n\b"
    ],
    "Huacho Díaz Mena": [
        r"\bJoaqu[ií]n\s*Huacho\s*D[ií]az\s*Mena\b",
        r"\bHuacho\s*D[ií]az\s*Mena\b"
    ]
}

# ----------------------------
# SESSION STATE
# ----------------------------

if "selected_pages" not in st.session_state:
    st.session_state.selected_pages = {}

# Estructura:
# {
#   unique_key: {
#       "file_name": "...",
#       "page_number": 1,
#       "png_bytes": b"..."
#   }
# }

# ----------------------------
# FUNCIONES AUXILIARES
# ----------------------------

def normalize_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = text.replace("\r", "\n")
    text = text.replace("\t", " ")

    # Normalizar espacios raros, pero aún conservar saltos
    text = re.sub(r"[ \f\v]+", " ", text)

    return text.strip()


def limpiar_texto_para_busqueda(text: str) -> str:
    """
    Limpieza estricta:
    1) une palabras partidas por guión + salto de línea
    2) elimina guiones restantes
    3) deja solo:
       - letras
       - números
       - espacios
       - vocales con acento
       - ñ / Ñ
    """
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = text.replace("\r", "\n")

    # Repetir hasta que ya no haya uniones pendientes por OCR raro
    prev = None
    while prev != text:
        prev = text
        text = re.sub(
            r"([A-Za-zÁÉÍÓÚáéíóúÑñ])\-\s*\n\s*([A-Za-zÁÉÍÓÚáéíóúÑñ])",
            r"\1\2",
            text
        )

    # Convertir saltos de línea a espacio
    text = text.replace("\n", " ")

    # Eliminar guiones restantes completamente
    text = text.replace("-", "")

    # Dejar solo letras, números, espacios, acentos y ñ
    text = re.sub(r"[^0-9A-Za-zÁÉÍÓÚáéíóúÑñ ]+", " ", text)

    # Colapsar espacios
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def normalizar_para_busqueda(text: str) -> str:
    """
    Normaliza solo para búsqueda interna:
    - minúsculas
    - quita acentos
    - conserva ñ
    """
    if not text:
        return ""

    text = text.lower().strip()

    reemplazos = {
        "á": "a",
        "é": "e",
        "í": "i",
        "ó": "o",
        "ú": "u",
        "Á": "a",
        "É": "e",
        "Í": "i",
        "Ó": "o",
        "Ú": "u",
    }

    for a, b in reemplazos.items():
        text = text.replace(a, b)

    text = re.sub(r"\s+", " ", text)
    return text.strip()


def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    gray = ImageOps.autocontrast(gray)
    gray = gray.filter(ImageFilter.SHARPEN)
    bw = gray.point(lambda x: 0 if x < 180 else 255, "1")
    return bw.convert("L")


def build_normalized_mapping(original_text: str):
    """
    Crea el texto normalizado y el mapa de índices
    para poder buscar en el normalizado y recortar
    snippets del original.
    """
    normalized_chars = []
    index_map = []

    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u",
        "Á": "a", "É": "e", "Í": "i", "Ó": "o", "Ú": "u",
    }

    for idx, ch in enumerate(original_text):
        ch2 = replacements.get(ch, ch)
        ch2 = ch2.lower()
        normalized_chars.append(ch2)
        index_map.append(idx)

    return "".join(normalized_chars), index_map


def make_unique_page_key(file_name: str, page_number: int) -> str:
    raw = f"{file_name}__{page_number}"
    return hashlib.md5(raw.encode("utf-8")).hexdigest()


# ----------------------------
# OCR SIN TIMEOUT
# ----------------------------

def hacer_ocr(img):
    try:
        texto = pytesseract.image_to_string(
            img,
            lang=OCR_LANG,
            config="--oem 3 --psm 6"
        )
        return texto or "", None
    except Exception as e:
        return "", str(e)


# ----------------------------
# RENDER DE PÁGINA
# ----------------------------

def render_page_png_bytes(page, dpi=DISPLAY_DPI) -> bytes:
    pix = page.get_pixmap(dpi=dpi, alpha=False)
    return pix.tobytes("png")


# ----------------------------
# EXTRAER TEXTO
# ----------------------------

def extract_text(page, page_number=None, debug_expander=None):
    """
    Devuelve:
    - text_original: limpio para mostrar snippets
    - text_search: normalizado para búsqueda
    - source: EMBEDDED / OCR / ERROR
    """
    texto_raw = ""

    # 1) Intentar texto embebido
    try:
        texto_raw = page.get_text("text") or ""
    except Exception as e:
        texto_raw = ""
        if DEBUG_MODE and debug_expander:
            debug_expander.write(f"Error extrayendo texto embebido: {e}")

    # 2) Si es muy poco, intentar bloques
    if len(texto_raw.strip()) < MIN_CHARS_TEXTO_EMBEDIDO:
        try:
            bloques = page.get_text("blocks") or []
            texto_bloques = "\n".join(
                b[4].strip() for b in bloques
                if len(b) >= 5 and str(b[4]).strip()
            )
            if len(texto_bloques.strip()) > len(texto_raw.strip()):
                texto_raw = texto_bloques
        except Exception as e:
            if DEBUG_MODE and debug_expander:
                debug_expander.write(f"No se pudo extraer texto por bloques: {e}")

    texto_original = limpiar_texto_para_busqueda(normalize_text(texto_raw))
    texto_busqueda = normalizar_para_busqueda(texto_original)

    if DEBUG_MODE and debug_expander:
        try:
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            image_blocks = sum(1 for b in blocks if b.get("type") == 1)
            text_blocks = sum(1 for b in blocks if b.get("type") == 0)

            debug_expander.write(f"Texto embebido limpio: {len(texto_original)} caracteres")
            debug_expander.write(f"Texto búsqueda normalizado: {len(texto_busqueda)} caracteres")
            debug_expander.write(
                f"Bloques totales: {len(blocks)} | "
                f"bloques texto: {text_blocks} | "
                f"bloques imagen: {image_blocks}"
            )
        except Exception as e:
            debug_expander.write(f"No se pudo leer page.get_text('dict'): {e}")

    # 3) Si el texto embebido sirve, usarlo
    if len(texto_original) >= MIN_CHARS_TEXTO_EMBEDIDO:
        if DEBUG_MODE and debug_expander:
            debug_expander.write("Se usará texto embebido limpio.")
            if texto_original:
                debug_expander.text_area(
                    f"Preview texto original limpio página {page_number}",
                    texto_original[:2000],
                    height=220
                )

        return {
            "text_original": texto_original,
            "text_search": texto_busqueda,
            "source": "EMBEDDED"
        }

    # 4) Si no, OCR
    if DEBUG_MODE and debug_expander:
        debug_expander.write("Se usará OCR.")

    try:
        pix = page.get_pixmap(
            dpi=OCR_DPI,
            colorspace=fitz.csGRAY,
            alpha=False
        )

        img = Image.frombytes(
            "L",
            [pix.width, pix.height],
            pix.samples
        )

        img = preprocess_for_ocr(img)

        text_ocr_raw, ocr_error = hacer_ocr(img)
        text_ocr_original = limpiar_texto_para_busqueda(normalize_text(text_ocr_raw))
        text_ocr_search = normalizar_para_busqueda(text_ocr_original)

        if DEBUG_MODE and debug_expander:
            debug_expander.write(f"OCR limpio caracteres: {len(text_ocr_original)}")
            debug_expander.write(f"Error OCR: {ocr_error}")
            if text_ocr_original:
                debug_expander.text_area(
                    f"Preview OCR original limpio página {page_number}",
                    text_ocr_original[:2000],
                    height=220
                )

        del pix, img
        gc.collect()

        return {
            "text_original": text_ocr_original,
            "text_search": text_ocr_search,
            "source": "OCR"
        }

    except Exception as e:
        if DEBUG_MODE and debug_expander:
            debug_expander.write(f"Fallo OCR en página {page_number}: {e}")

        return {
            "text_original": "",
            "text_search": "",
            "source": "ERROR"
        }


# ----------------------------
# BUSCAR PALABRAS
# ----------------------------

def search_keywords(text_original, text_search):
    """
    Busca sobre texto normalizado para tolerar acentos,
    pero recorta el snippet desde el texto original.
    """
    results = []
    seen = set()

    if not text_original or not text_search:
        return results

    normalized_original, index_map = build_normalized_mapping(text_original)

    for kw, patterns in KEYWORDS.items():
        for pattern in patterns:
            pattern_norm = normalizar_para_busqueda(pattern)

            pattern_norm = pattern_norm.replace(r"[aá]", r"[aa]")
            pattern_norm = pattern_norm.replace(r"[eé]", r"[ee]")
            pattern_norm = pattern_norm.replace(r"[ií]", r"[ii]")
            pattern_norm = pattern_norm.replace(r"[oó]", r"[oo]")
            pattern_norm = pattern_norm.replace(r"[uú]", r"[uu]")

            try:
                matches = list(re.finditer(pattern_norm, normalized_original, re.IGNORECASE))
            except re.error:
                continue

            for match in matches:
                start_norm = max(match.start() - 60, 0)
                end_norm = min(match.end() + 60, len(normalized_original))

                start_orig = index_map[start_norm]
                end_orig = index_map[end_norm - 1] + 1 if end_norm > start_norm else index_map[start_norm] + 1

                phrase = text_original[start_orig:end_orig]
                phrase = re.sub(r"\s+", " ", phrase).strip()

                phrase_key = normalizar_para_busqueda(phrase)
                if phrase_key in seen:
                    continue
                seen.add(phrase_key)

                highlighted = phrase
                try:
                    highlighted = re.sub(
                        pattern,
                        r"<span style='color:red;font-weight:bold'>\g<0></span>",
                        phrase,
                        flags=re.IGNORECASE
                    )
                except re.error:
                    pass

                results.append(highlighted)

    return results


# ----------------------------
# DOCX
# ----------------------------

def build_docx_from_selected_pages(selected_pages_dict: dict) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    items = sorted(
        selected_pages_dict.values(),
        key=lambda x: (x["file_name"].lower(), x["page_number"])
    )

    for idx, item in enumerate(items):
        if idx > 0:
            doc.add_page_break()

        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(f'{item["file_name"]} - Página {item["page_number"]}')
        run1.bold = True

        doc.add_paragraph("")

        img_stream = io.BytesIO(item["png_bytes"])
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run()
        run2.add_picture(img_stream, width=Inches(7.2))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ----------------------------
# UI CSS
# ----------------------------

st.markdown("""
<style>
.small-preview-card {
    border: 1px solid #DDD;
    border-radius: 10px;
    padding: 10px;
    margin-bottom: 14px;
    background: #FFF;
}
.small-preview-title {
    font-size: 15px;
    font-weight: 600;
    margin-bottom: 8px;
}
.small-preview-meta {
    font-size: 12px;
    color: #666;
    margin-bottom: 8px;
}
</style>
""", unsafe_allow_html=True)

# ----------------------------
# UI
# ----------------------------

st.markdown("<h1 style='text-align:center'>Noticias Yucatán</h1>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align:center;color:gray'>Análisis automático de PDFs</h3>", unsafe_allow_html=True)

st.write("---")

uploaded_files = st.file_uploader(
    "Sube tus PDFs",
    type=["pdf"],
    accept_multiple_files=True
)

top_col1, top_col2 = st.columns([2, 1])

with top_col1:
    st.write(f"Páginas seleccionadas para DOCX: **{len(st.session_state.selected_pages)}**")

with top_col2:
    if st.session_state.selected_pages:
        docx_bytes = build_docx_from_selected_pages(st.session_state.selected_pages)
        st.download_button(
            label="Descargar DOCX de páginas seleccionadas",
            data=docx_bytes,
            file_name="paginas_seleccionadas_yucatan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

if st.button("Limpiar páginas seleccionadas"):
    st.session_state.selected_pages = {}
    st.rerun()

summary = {}

# ----------------------------
# PROCESAMIENTO
# ----------------------------

if uploaded_files:
    for file in uploaded_files:
        st.header(file.name)
        summary[file.name] = []

        try:
            pdf_bytes = file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        except Exception:
            st.warning("No se pudo abrir el PDF")
            continue

        total = len(doc)
        progress = st.progress(0)

        for i, page in enumerate(doc):
            page_number = i + 1
            progress.progress((i + 1) / total)

            debug_expander = None
            if DEBUG_MODE:
                debug_expander = st.expander(f"Diagnóstico página {page_number}", expanded=False)

            extracted = extract_text(page, page_number=page_number, debug_expander=debug_expander)
            text_original = extracted["text_original"]
            text_search = extracted["text_search"]

            results = search_keywords(text_original, text_search)

            if DEBUG_MODE and debug_expander:
                debug_expander.write(f"Fuente usada: {extracted['source']}")
                debug_expander.write(f"Coincidencias encontradas: {len(results)}")

            if results:
                page_png_bytes = render_page_png_bytes(page, dpi=DISPLAY_DPI)
                unique_key = make_unique_page_key(file.name, page_number)
                already_selected = unique_key in st.session_state.selected_pages

                col1, col2 = st.columns([1, 3])

                with col1:
                    st.markdown(
                        f"""
                        <div class="small-preview-card">
                            <div class="small-preview-title">Página {page_number}</div>
                            <div class="small-preview-meta">Vista en alta calidad reducida visualmente</div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    btn_label = "Agregada al DOCX" if already_selected else "Agregar página al DOCX"
                    if st.button(
                        btn_label,
                        key=f"add_docx_{unique_key}",
                        use_container_width=True,
                        disabled=already_selected
                    ):
                        st.session_state.selected_pages[unique_key] = {
                            "file_name": file.name,
                            "page_number": page_number,
                            "png_bytes": page_png_bytes
                        }
                        st.rerun()

                    st.download_button(
                        label="Descargar página PNG",
                        data=page_png_bytes,
                        file_name=f"{file.name}_pagina_{page_number}.png",
                        mime="image/png",
                        key=f"download_png_{unique_key}",
                        use_container_width=True
                    )

                    st.image(
                        page_png_bytes,
                        width=DISPLAY_WIDTH
                    )

                with col2:
                    st.markdown(f"### Página {page_number}")

                    for r in results:
                        st.markdown(r, unsafe_allow_html=True)

                    st.caption(f"Fuente de texto usada: {extracted['source']}")

                summary[file.name].append(
                    f"Página {page_number}: " + re.sub("<[^<]+?>", "", results[0])
                )

            gc.collect()

        progress.progress(1.0)

        if summary[file.name]:
            st.success("Procesamiento terminado con coincidencias")
        else:
            st.info("No se encontraron coincidencias")

        doc.close()

# ----------------------------
# RESUMEN FINAL
# ----------------------------

final = ""

for pdf, entries in summary.items():
    final += pdf + "\n"

    if entries:
        final += "\n".join(entries)
    else:
        final += "Sin coincidencias"

    final += "\n\n"

if final:
    st.text_area(
        "Resumen final (copiar)",
        final,
        height=400
    )